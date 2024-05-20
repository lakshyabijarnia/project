import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from scipy.signal import find_peaks, welch
from docx import Document
from docx.shared import Inches

class GLevelPSDApp(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("G-Level and PSD Plotter")
        self.geometry("1000x800")

        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill='both', expand=True)

        self.input_tab = ttk.Frame(self.notebook)
        self.glevel_tab = ttk.Frame(self.notebook)
        self.psd_tab = ttk.Frame(self.notebook)

        self.notebook.add(self.input_tab, text='Inputs')
        self.notebook.add(self.glevel_tab, text='G-Levels')
        self.notebook.add(self.psd_tab, text='PSD Plots')

        self.create_input_tab()
        self.create_glevel_tab()
        self.create_psd_tab()

        self.data = None
        self.velocity_data = None
        self.sensitivity = None
        self.sampling_freq = None
        self.selected_range = None

    def create_input_tab(self):
        ttk.Label(self.input_tab, text="Sensor Sensitivity:").grid(row=0, column=0, padx=10, pady=10)
        self.sensitivity_entry = ttk.Entry(self.input_tab)
        self.sensitivity_entry.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(self.input_tab, text="Sampling Frequency:").grid(row=1, column=0, padx=10, pady=10)
        self.sampling_freq_entry = ttk.Entry(self.input_tab)
        self.sampling_freq_entry.grid(row=1, column=1, padx=10, pady=10)

        ttk.Button(self.input_tab, text="Load CSV/Excel File", command=self.load_file).grid(row=2, column=0, columnspan=2, padx=10, pady=10)
        ttk.Button(self.input_tab, text="Load Velocity Profile", command=self.load_velocity_profile).grid(row=3, column=0, columnspan=2, padx=10, pady=10)

        ttk.Button(self.input_tab, text="Plot G-Levels", command=self.plot_glevels).grid(row=4, column=0, columnspan=2, padx=10, pady=10)

    def create_glevel_tab(self):
        self.glevel_canvas_frame = tk.Canvas(self.glevel_tab)
        self.scrollbar = ttk.Scrollbar(self.glevel_tab, orient="vertical", command=self.glevel_canvas_frame.yview)
        self.glevel_canvas = ttk.Frame(self.glevel_canvas_frame)

        self.glevel_canvas_frame.create_window((0, 0), window=self.glevel_canvas, anchor="nw")
        self.glevel_canvas_frame.configure(yscrollcommand=self.scrollbar.set)

        self.glevel_canvas_frame.pack(side=tk.LEFT, fill='both', expand=True)
        self.scrollbar.pack(side=tk.RIGHT, fill='y')

        self.glevel_canvas.bind("<Configure>", lambda e: self.glevel_canvas_frame.configure(scrollregion=self.glevel_canvas_frame.bbox("all")))

        self.glevel_plots = []
        self.glevel_figs = []
        self.glevel_axs = []

        ttk.Button(self.glevel_tab, text="Plot PSD from Selection", command=self.plot_psd_from_selection).pack(padx=10, pady=10)
        ttk.Button(self.glevel_tab, text="Export Plots to DOCX", command=self.export_plots_to_docx).pack(padx=10, pady=10)

    def create_psd_tab(self):
        self.psd_canvas_frame = tk.Canvas(self.psd_tab)
        self.scrollbar_psd = ttk.Scrollbar(self.psd_tab, orient="vertical", command=self.psd_canvas_frame.yview)
        self.psd_canvas = ttk.Frame(self.psd_canvas_frame)

        self.psd_canvas_frame.create_window((0, 0), window=self.psd_canvas, anchor="nw")
        self.psd_canvas_frame.configure(yscrollcommand=self.scrollbar_psd.set)

        self.psd_canvas_frame.pack(side=tk.LEFT, fill='both', expand=True)
        self.scrollbar_psd.pack(side=tk.RIGHT, fill='y')

        self.psd_canvas.bind("<Configure>", lambda e: self.psd_canvas_frame.configure(scrollregion=self.psd_canvas_frame.bbox("all")))

        self.psd_plots = []
        self.psd_figs = []
        self.psd_axs = []

    def load_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv"), ("Excel files", "*.xlsx")])
        if file_path:
            if file_path.endswith('.csv'):
                self.data = pd.read_csv(file_path)
            elif file_path.endswith('.xlsx'):
                self.data = pd.read_excel(file_path)
            messagebox.showinfo("File Loaded", "Vibration profile loaded successfully.")
            
    def load_velocity_profile(self):
        file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv"), ("Excel files", "*.xlsx")])
        if file_path:
            if file_path.endswith('.csv'):
                self.velocity_data = pd.read_csv(file_path)
            elif file_path.endswith('.xlsx'):
                self.velocity_data = pd.read_excel(file_path)
            messagebox.showinfo("File Loaded", "Velocity profile loaded successfully.")
            
    def plot_glevels(self):
        if self.data is not None:
            try:
                self.sensitivity = float(self.sensitivity_entry.get())
                self.sampling_freq = float(self.sampling_freq_entry.get())
            except ValueError:
                messagebox.showerror("Input Error", "Please enter valid numbers for sensitivity and sampling frequency.")
                return

            time = self.data.iloc[:, 0]
            channel_data = self.data.iloc[:, 1:] * self.sensitivity

            for plot in self.glevel_plots:
                plot.get_tk_widget().pack_forget()
            self.glevel_plots.clear()
            self.glevel_figs.clear()
            self.glevel_axs.clear()

            for i in range(0, min(channel_data.shape[1], 24), 3):
                fig, axs = plt.subplots(3, 1, figsize=(8, 8))
                self.glevel_figs.append(fig)

                for j, ax in enumerate(axs):
                    ax.clear()
                    ax.plot(time, channel_data.iloc[:, i+j], label=f'Channel {i//3 + 1} - {"XYZ"[j]}')
                    ax.set_xlabel("Time")
                    ax.set_ylabel("G-Levels")
                    ax.legend()

                    if self.velocity_data is not None:
                        ax2 = ax.twinx()
                        velocity_time = self.velocity_data.iloc[:, 0]
                        velocity = self.velocity_data.iloc[:, 1]
                        ax2.plot(velocity_time, velocity, label='Velocity', linestyle='--', color='orange')
                        ax2.set_ylabel("Velocity")
                        ax2.legend(loc='upper right')

                self.glevel_canvas.update_idletasks()
                plot = FigureCanvasTkAgg(fig, master=self.glevel_canvas)
                plot.get_tk_widget().pack(fill='both', expand=True)
                self.glevel_plots.append(plot)

            # Create toolbar outside the loop
            if self.glevel_plots:
                plot = self.glevel_plots[0]
                plot.toolbar = NavigationToolbar2Tk(plot, self.glevel_canvas)
                plot.toolbar.update()
                plot.get_tk_widget().pack(fill='both', expand=True)

            self.notebook.select(self.glevel_tab)
        else:
            messagebox.showerror("Data Error", "Please load the data file first.")

    def plot_psd_from_selection(self):
        if self.data is not None:
            for plot in self.psd_plots:
                plot.get_tk_widget().pack_forget()
            self.psd_plots.clear()
            self.psd_figs.clear()
            self.psd_axs.clear()

            for i in range(0, min(self.data.shape[1]-1, 24), 3):
                fig, axs = plt.subplots(3, 1, figsize=(8, 8))
                self.psd_figs.append(fig)

                for j, ax in enumerate(axs):
                    if self.selected_range is not None:
                        start, end = self.selected_range
                        data_subset = self.data[(self.data.iloc[:, 0] >= start) & (self.data.iloc[:, 0] <= end)]
                    else:
                        data_subset = self.data

                    time = data_subset.iloc[:, 0]
                    channel_data = data_subset.iloc[:, i+1] * self.sensitivity

                    f, Pxx = welch(channel_data, fs=self.sampling_freq)

                    ax.clear()
                    ax.semilogy(f, Pxx, label=f'Channel {i//3 + 1} - {"XYZ"[j]}')
                    ax.set_xlabel("Frequency [Hz], Log Scale")
                    ax.set_ylabel("PSD [V**2/Hz], Log Scale")
                    ax.set_yscale('log')
                    ax.set_xscale('log')
                    ax.legend()

                self.psd_canvas.update_idletasks()
                plot = FigureCanvasTkAgg(fig, master=self.psd_canvas)
                plot.get_tk_widget().pack(fill='both', expand=True)
                self.psd_plots.append(plot)

            # Create toolbar outside the loop
            if self.psd_plots:
                plot = self.psd_plots[0]
                plot.toolbar = NavigationToolbar2Tk(plot, self.psd_canvas)
                plot.toolbar.update()
                plot.get_tk_widget().pack(fill='both', expand=True)

            self.notebook.select(self.psd_tab)
        else:
            messagebox.showerror("Data Error", "Please load the data file first.")

    def show_peaks(self):
        for fig in self.glevel_figs:
            for ax in fig.axes:
                for line in ax.get_lines():
                    xdata = line.get_xdata()
                    ydata = line.get_ydata()
                    peaks, _ = find_peaks(ydata)
                    peak_values = [(xdata[idx], ydata[idx]) for idx in peaks]
                    ax.plot(*zip(*peak_values), marker='o', markersize=8, color='red', linestyle='None')
                    for peak in peak_values:
                        ax.text(peak[0], peak[1], f'({peak[0]:.2f}, {peak[1]:.2f})', fontsize=8, color='red')

    def export_plots_to_docx(self):
        if not self.glevel_figs and not self.psd_figs:
            messagebox.showerror("Export Error", "No plots available to export.")
            return

        doc = Document()
        doc.add_heading('G-Level and PSD Plots', level=1)
 
        if self.glevel_figs:
            doc.add_heading('G-Level Plots', level=2)
            for i, fig in enumerate(self.glevel_figs):
                fig.savefig(f'glevel_plot_{i}.png')
                doc.add_picture(f'glevel_plot_{i}.png', width=Inches(6))

        if self.psd_figs:
            doc.add_heading('PSD Plots', level=2)
            for i, fig in enumerate(self.psd_figs):
                fig.savefig(f'psd_plot_{i}.png')
                doc.add_picture(f'psd_plot_{i}.png', width=Inches(6))

        doc.save('plots.docx')
        messagebox.showinfo("Export Success", "Plots have been exported to plots.docx")

if __name__ == "__main__":
    app = GLevelPSDApp()
    app.mainloop()
