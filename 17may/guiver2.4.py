import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
import concurrent.futures
from scipy.signal import welch
import os

class GLevelPSDApp(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("G-Level and PSD Plotter")
        self.geometry("1000x800")

        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill='both', expand=True)

        self.input_tab = ttk.Frame(self.notebook)
        self.glevel_tab = ttk.Frame(self.notebook)
        self.psd_tab = ttk.Frame(self.notebook)

        self.notebook.add(self.input_tab, text='Inputs')
        self.notebook.add(self.glevel_tab, text='G-Levels')
        self.notebook.add(self.psd_tab, text='PSD Plots')

        self.data = None
        self.velocity_data = None
        self.sensitivity = None
        self.sampling_freq = None
        self.selected_range = None
        self.velocity_present = tk.BooleanVar()

        self.create_input_tab()
        self.create_glevel_tab()
        self.create_psd_tab()

    def create_input_tab(self):
        ttk.Label(self.input_tab, text="Sensor Sensitivity:").grid(row=0, column=0, padx=10, pady=10)
        self.sensitivity_entry = ttk.Entry(self.input_tab)
        self.sensitivity_entry.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(self.input_tab, text="Sampling Frequency:").grid(row=1, column=0, padx=10, pady=10)
        self.sampling_freq_entry = ttk.Entry(self.input_tab)
        self.sampling_freq_entry.grid(row=1, column=1, padx=10, pady=10)

        self.velocity_checkbox = ttk.Checkbutton(self.input_tab, text="Velocity Data Present", variable=self.velocity_present, command=self.toggle_velocity_profile)
        self.velocity_checkbox.grid(row=2, column=0, columnspan=2, padx=10, pady=10)

        self.load_velocity_button = ttk.Button(self.input_tab, text="Load Velocity Profile", command=self.load_velocity_profile, state=tk.DISABLED)
        self.load_velocity_button.grid(row=3, column=0, columnspan=2, padx=10, pady=10)

        ttk.Label(self.input_tab, text="Number of Files to Upload:").grid(row=4, column=0, padx=10, pady=10)
        self.num_files_spinbox = ttk.Spinbox(self.input_tab, from_=1, to=10)
        self.num_files_spinbox.grid(row=4, column=1, padx=10, pady=10)

        self.upload_files_button = ttk.Button(self.input_tab, text="Upload File(s)", command=self.upload_files)
        self.upload_files_button.grid(row=5, column=0, columnspan=2, padx=10, pady=10)

        self.uploaded_files_label = ttk.Label(self.input_tab, text="")
        self.uploaded_files_label.grid(row=6, column=0, columnspan=2, padx=10, pady=10)

        ttk.Button(self.input_tab, text="Plot G-Levels", command=self.plot_glevels).grid(row=7, column=0, columnspan=2, padx=10, pady=10)
        ttk.Button(self.input_tab, text="Plot PSD", command=self.plot_psd_from_selection).grid(row=8, column=0, columnspan=2, padx=10, pady=10)
        ttk.Button(self.input_tab, text="Export Plots", command=self.export_plots).grid(row=9, column=0, columnspan=2, padx=10, pady=10)

    def toggle_velocity_profile(self):
        if self.velocity_present.get():
            self.load_velocity_button.config(state=tk.NORMAL)
        else:
            self.load_velocity_button.config(state=tk.DISABLED)
            self.velocity_data = None  # Clear velocity data if checkbox is unchecked

    def create_glevel_tab(self):
        self.glevel_canvas_frame = tk.Canvas(self.glevel_tab)
        self.glevel_canvas_frame.pack(side=tk.LEFT, fill='both', expand=True)
        self.scrollbar = ttk.Scrollbar(self.glevel_tab, orient="vertical", command=self.glevel_canvas_frame.yview)
        self.scrollbar.pack(side=tk.RIGHT, fill='y')
        self.glevel_canvas = ttk.Frame(self.glevel_canvas_frame)

        self.glevel_canvas_frame.create_window((0, 0), window=self.glevel_canvas, anchor="nw")
        self.glevel_canvas_frame.configure(yscrollcommand=self.scrollbar.set)

        self.glevel_canvas.bind("<Configure>", lambda e: self.glevel_canvas_frame.configure(scrollregion=self.glevel_canvas_frame.bbox("all")))

        self.glevel_plots = []
        self.glevel_figs = []
        self.glevel_axs = []

    def create_psd_tab(self):
        self.psd_canvas_frame = tk.Canvas(self.psd_tab)
        self.psd_canvas_frame.pack(side=tk.LEFT, fill='both', expand=True)
        self.scrollbar_psd = ttk.Scrollbar(self.psd_tab, orient="vertical", command=self.psd_canvas_frame.yview)
        self.scrollbar_psd.pack(side=tk.RIGHT, fill='y')
        self.psd_canvas = ttk.Frame(self.psd_canvas_frame)

        self.psd_canvas_frame.create_window((0, 0), window=self.psd_canvas, anchor="nw")
        self.psd_canvas_frame.configure(yscrollcommand=self.scrollbar_psd.set)

        self.psd_canvas.bind("<Configure>", lambda e: self.psd_canvas_frame.configure(scrollregion=self.psd_canvas_frame.bbox("all")))

        self.psd_plots = []
        self.psd_figs = []
        self.psd_axs = []


    def upload_files(self):
        num_files = int(self.num_files_spinbox.get())
        self.uploaded_files = []
        for _ in range(num_files):
            file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv"), ("Excel files", "*.xls"), ("Excel files", "*.xlsx")])
            if file_path:
                self.uploaded_files.append(file_path)
        self.uploaded_files_label.config(text=", ".join([os.path.basename(file) for file in self.uploaded_files]))
        self.load_data_parallel(self.uploaded_files)

    def load_data_parallel(self, file_paths):
        dfs = []
        for file_path in file_paths:
            if file_path.endswith('.xlsx'):
                df = pd.read_excel(file_path, engine='openpyxl')
            elif file_path.endswith('.xls'):
                df = pd.read_excel(file_path)
            elif file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                raise ValueError("Unsupported file format")
            dfs.append(df)
    
        # Extract time column from the first file
        time_column = dfs[0].iloc[:, 0]
    
        # Create a dictionary to store data frames for each sensor
        sensor_dfs = {}
    
        # Iterate over the data frames and separate data for each sensor
        for i in range(len(dfs) // 3):
            sensor_name = f"Sensor {i+1}"
            sensor_dfs[sensor_name] = dfs[i * 3: (i + 1) * 3]
    
        # Merge data frames for each sensor
        merged_dfs = {}
        for sensor_name, dfs in sensor_dfs.items():
            merged_df = pd.concat(dfs, axis=1)
            merged_dfs[sensor_name] = merged_df
    
        return time_column, merged_dfs
    
    







    # def load_file(self):
    #     for file_path in self.uploaded_files:
    #         # file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv"), ("Excel files", "*.xlsx")])
    #         if file_path:
    #             try:
    #                 if file_path.endswith('.xlsx'):
    #                     self.data = self.load_excel_parallel(file_path)
    #                 elif file_path.endswith('.csv'):
    #                     self.data = self.load_csv_parallel(file_path)
    #                 messagebox.showinfo("File Loaded", "Vibration profile loaded successfully.")
    #             except Exception as e:
    #                 messagebox.showerror("File Error", f"An error occurred while loading the file: {e}")

    # def load_csv_parallel(self, file_path):
    #     with concurrent.futures.ThreadPoolExecutor() as executor:
    #         df = pd.read_csv(file_path, header='infer')
    #     return df

    # def load_excel_parallel(self, file_path):
    #     with concurrent.futures.ThreadPoolExecutor() as executor:
    #         df = pd.read_excel(file_path, header='infer')
    #     return df

    def load_velocity_profile(self):
        file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv"), ("Excel files", "*.xlsx")])
        if file_path:
            try:
                if file_path.endswith('.csv'):
                    self.velocity_data = self.load_csv_parallel(file_path)
                elif file_path.endswith('.xlsx'):
                    self.velocity_data = self.load_excel_parallel(file_path)
                messagebox.showinfo("File Loaded", "Velocity profile loaded successfully.")
            except Exception as e:
                messagebox.showerror("File Error", f"An error occurred while loading the file: {e}")

    def plot_glevels(self):
        if self.data is not None:
            try:
                self.sensitivity = float(self.sensitivity_entry.get())
                self.sampling_freq = float(self.sampling_freq_entry.get())
            except ValueError:
                messagebox.showerror("Input Error", "Please enter valid numbers for sensitivity and sampling frequency.")
                return

            self.clear_plots(self.glevel_plots, self.glevel_figs, self.glevel_axs)

            time = self.data.iloc[:, 0]
            channel_data = self.data.iloc[:, 1:] * self.sensitivity

            for i in range(0, min(channel_data.shape[1], 24), 3):
                fig, axs = plt.subplots(3, 1, figsize=(10, 8))
                self.glevel_figs.append(fig)

                for j, ax in enumerate(axs):
                    ax.clear()
                    ax.plot(time, channel_data.iloc[:, i + j], label=f'Channel {i // 3 + 1} - {"XYZ"[j]}')
                    ax.set_xlabel("Time")
                    ax.set_ylabel("G-Levels")
                    ax.legend(loc='upper right')

                    if self.velocity_data is not None and self.velocity_present.get():
                        velocity_time = self.velocity_data.iloc[:, 0]
                        velocity = self.velocity_data.iloc[:, 1]
                        ax_velocity = ax.twinx()
                        ax_velocity.plot(velocity_time, velocity, label='Velocity', linestyle='--', color='red')
                        ax_velocity.set_ylabel("Velocity")
                        ax_velocity.legend(loc='upper left')

                self.glevel_canvas.update_idletasks()
                plot = FigureCanvasTkAgg(fig, master=self.glevel_canvas)
                plot.get_tk_widget().pack(fill='both', expand=True)
                plot.toolbar = NavigationToolbar2Tk(plot, self.glevel_canvas)
                plot.toolbar.update()
                self.glevel_plots.append(plot)

            self.notebook.select(self.glevel_tab)
        else:
            messagebox.showerror("Data Error", "Please load the data file first.")

    def plot_psd_from_selection(self):
        if self.data is not None:
            self.clear_plots(self.psd_plots, self.psd_figs, self.psd_axs)
            selected_channel = 0  # Default to first channel if no range is selected

            if self.selected_range is not None:
                selected_channel = int(self.selected_range[0]) // 3

            time = self.data.iloc[:, 0]
            data = self.data.iloc[:, 1 + selected_channel * 3: 1 + (selected_channel + 1) * 3] * self.sensitivity

            fig, axs = plt.subplots(3, 1, figsize=(10, 8))
            self.psd_figs.append(fig)

            for i, ax in enumerate(axs):
                f, Pxx = welch(data.iloc[:, i], fs=self.sampling_freq, nperseg=1024)
                ax.clear()
                ax.semilogy(f, np.sqrt(Pxx), label=f'Channel {selected_channel + 1} - {"XYZ"[i]}')
                ax.set_xlabel("Frequency (Hz)")
                ax.set_ylabel("PSD (G^2/Hz)")
                ax.legend(loc='upper right')

            self.psd_canvas.update_idletasks()
            plot = FigureCanvasTkAgg(fig, master=self.psd_canvas)
            plot.get_tk_widget().pack(fill='both', expand=True)
            plot.toolbar = NavigationToolbar2Tk(plot, self.psd_canvas)
            plot.toolbar.update()
            self.psd_plots.append(plot)
        
        self.notebook.select(self.psd_tab)

    def export_plots(self):
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading('G-Levels and PSD Plots', 0)

        for fig in self.glevel_figs:
            img_path = '/tmp/glevel_plot.png'
            fig.savefig(img_path)
            doc.add_paragraph('G-Level Plot')
            doc.add_picture(img_path, width=Inches(6))

        for fig in self.psd_figs:
            img_path = '/tmp/psd_plot.png'
            fig.savefig(img_path)
            doc.add_paragraph('PSD Plot')
            doc.add_picture(img_path, width=Inches(6))

        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx")])
        if save_path:
            doc.save(save_path)
            messagebox.showinfo("Export Successful", "Plots exported successfully.")
    def clear_plots(self, plots, figs, axs):
        for plot in plots:
            plot.get_tk_widget().destroy()
        plots.clear()
        figs.clear()
        axs.clear()

if __name__ == "__main__":
    app = GLevelPSDApp()
    app.mainloop()

