import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from scipy.signal import find_peaks, welch

class GLevelPSDApp(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("G-Level and PSD Plotter")
        self.geometry("1000x800")

        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill='both', expand=True)

        self.input_tab = ttk.Frame(self.notebook)
        self.glevel_tab = ttk.Frame(self.notebook)
        self.psd_tab = ttk.Frame(self.notebook)

        self.notebook.add(self.input_tab, text='Inputs')
        self.notebook.add(self.glevel_tab, text='G-Levels')
        self.notebook.add(self.psd_tab, text='PSD Plots')

        self.data = None
        self.sensitivity = None
        self.sampling_freq = None
        self.selected_range = None

        self.create_input_tab()
        self.create_glevel_tab()
        self.create_psd_tab()

    def create_input_tab(self):
        ttk.Label(self.input_tab, text="Sensor Sensitivity (in mV):").grid(row=0, column=0, padx=10, pady=10)
        self.sensitivity_entry = ttk.Entry(self.input_tab)
        self.sensitivity_entry.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(self.input_tab, text="Sampling Frequency:").grid(row=1, column=0, padx=10, pady=10)
        self.sampling_freq_entry = ttk.Entry(self.input_tab)
        self.sampling_freq_entry.grid(row=1, column=1, padx=10, pady=10)

        ttk.Button(self.input_tab, text="Load CSV/Excel File", command=self.load_file).grid(row=2, column=0, columnspan=2, padx=10, pady=10)
        ttk.Button(self.input_tab, text="Plot G-Levels", command=self.plot_glevels).grid(row=3, column=0, columnspan=2, padx=10, pady=10)
        ttk.Button(self.input_tab, text="Plot PSD", command=self.plot_psd_from_selection).grid(row=4, column=0, columnspan=2, padx=10, pady=10)
        ttk.Button(self.input_tab, text="Export Plots", command=self.export_plots).grid(row=5, column=0, columnspan=2, padx=10, pady=10)

    def create_glevel_tab(self):
        self.glevel_canvas_frame = tk.Canvas(self.glevel_tab)
        self.glevel_canvas_frame.pack(side=tk.LEFT, fill='both', expand=True)
        self.scrollbar = ttk.Scrollbar(self.glevel_tab, orient="vertical", command=self.glevel_canvas_frame.yview)
        self.scrollbar.pack(side=tk.RIGHT, fill='y')
        self.glevel_canvas_frame.configure(yscrollcommand=self.scrollbar.set)

        self.glevel_canvas = ttk.Frame(self.glevel_canvas_frame)
        self.glevel_canvas_frame.create_window((0, 0), window=self.glevel_canvas, anchor="nw")
        self.glevel_canvas_frame.bind("<Configure>", lambda e: self.glevel_canvas_frame.configure(scrollregion=self.glevel_canvas_frame.bbox("all")))

        self.glevel_plots = []
        self.glevel_figs = []
        self.glevel_axs = []

    def create_psd_tab(self):
        self.psd_canvas_frame = tk.Canvas(self.psd_tab)
        self.psd_canvas_frame.pack(side=tk.LEFT, fill='both', expand=True)
        self.scrollbar_psd = ttk.Scrollbar(self.psd_tab, orient="vertical", command=self.psd_canvas_frame.yview)
        self.scrollbar_psd.pack(side=tk.RIGHT, fill='y')
        self.psd_canvas_frame.configure(yscrollcommand=self.scrollbar_psd.set)

        self.psd_canvas = ttk.Frame(self.psd_canvas_frame)
        self.psd_canvas_frame.create_window((0, 0), window=self.psd_canvas, anchor="nw")
        self.psd_canvas_frame.bind("<Configure>", lambda e: self.psd_canvas_frame.configure(scrollregion=self.psd_canvas_frame.bbox("all")))

        self.psd_plots = []
        self.psd_figs = []
        self.psd_axs = []

    def load_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv"), ("Excel files", "*.xlsx")])
        if file_path:
            if file_path.endswith('.csv'):
                self.data = pd.read_csv(file_path)
            elif file_path.endswith('.xlsx'):
                self.data = pd.read_excel(file_path)
            messagebox.showinfo("File Loaded", "Vibration profile loaded successfully.")

    def plot_glevels(self):
        if self.data is not None:
            try:
                self.sensitivity = float(self.sensitivity_entry.get()) / 1000  # Convert mV to V
                self.sampling_freq = float(self.sampling_freq_entry.get())
            except ValueError:
                messagebox.showerror("Input Error", "Please enter valid numbers for sensitivity and sampling frequency.")
                return

            time = self.data.iloc[:, 0]
            channel_data = self.data.iloc[:, 1:] * self.sensitivity

            for plot in self.glevel_plots:
                plot.get_tk_widget().pack_forget()
            self.glevel_plots.clear()
            self.glevel_figs.clear()
            self.glevel_axs.clear()

            for i in range(0, min(channel_data.shape[1], 24), 3):
                fig, axs = plt.subplots(3, 1, figsize=(10, 8))
                self.glevel_figs.append(fig)

                for j, ax in enumerate(axs):
                    ax.clear()
                    ax.plot(time, channel_data.iloc[:, i+j], label=f'Channel {i//3 + 1} - {"XYZ"[j]}')
                    ax.set_xlabel("Time")
                    ax.set_ylabel("G-Levels")
                    ax.legend(loc='upper right')

                    self.highlight_extreme_peaks(ax)

                plot = FigureCanvasTkAgg(fig, master=self.glevel_canvas)
                plot.get_tk_widget().pack(fill='both', expand=True)
                plot.toolbar = NavigationToolbar2Tk(plot, self.glevel_canvas)
                plot.toolbar.update()
                plot.get_tk_widget().pack(fill='both', expand=True)
                self.glevel_plots.append(plot)

            self.notebook.select(self.glevel_tab)
        else:
            messagebox.showerror("Data Error", "Please load the data file first.")

    def plot_psd_from_selection(self):
        if self.data is not None:
            for plot in self.psd_plots:
                plot.get_tk_widget().pack_forget()
            self.psd_plots.clear()
            self.psd_figs.clear()
            self.psd_axs.clear()

            for i in range(0, min(self.data.shape[1]-1, 24), 3):
                fig, axs = plt.subplots(3, 1, figsize=(10, 8))
                self.psd_figs.append(fig)

                for j, ax in enumerate(axs):
                    if self.selected_range is not None:
                        start, end = self.selected_range
                        data_subset = self.data[(self.data.iloc[:, 0] >= start) & (self.data.iloc[:, 0] <= end)]
                    else:
                        data_subset = self.data

                    time = data_subset.iloc[:, 0]
                    channel_data = data_subset.iloc[:, i + 1 + j] * self.sensitivity

                    f, Pxx = welch(channel_data, fs=self.sampling_freq)

                    ax.clear()
                    ax.semilogy(f, Pxx, label=f'Channel {i//3 + 1} - {"XYZ"[j]}')
                    ax.set_xscale("log")  # Set x-axis to logarithmic scale
                    ax.set_xlabel("Frequency [Hz]")
                    ax.set_ylabel("PSD [V**2/Hz]")
                    ax.legend(loc='upper right')

                    self.highlight_extreme_peaks(ax)

                plot = FigureCanvasTkAgg(fig, master=self.psd_canvas)
                plot.get_tk_widget().pack(fill='both', expand=True)
                plot.toolbar = NavigationToolbar2Tk(plot, self.psd_canvas)
                # plot.toolbar.update()
                plot.get_tk_widget().pack(fill='both', expand=True)
                self.psd_plots.append(plot)

            self.notebook.select(self.psd_tab)
        else:
            messagebox.showerror("Data Error", "Please load the data file first.")

    def highlight_extreme_peaks(self, ax):
        for line in ax.get_lines():
            xdata = line.get_xdata()
            ydata = line.get_ydata()
            if len(ydata) > 0:
                max_idx = np.argmax(ydata)
                min_idx = np.argmin(ydata)
                ax.plot(xdata[max_idx], ydata[max_idx], marker='o', markersize=8, color='red')
                ax.text(xdata[max_idx], ydata[max_idx], f'{ydata[max_idx]:}', fontsize=8, color='red', ha='left', va='bottom', bbox=dict(facecolor='white', alpha=0.5))
                ax.plot(xdata[min_idx], ydata[min_idx], marker='o', markersize=8, color='blue')
                ax.text(xdata[min_idx], ydata[min_idx], f'{ydata[min_idx]:}', fontsize=8, color='blue', ha='left', va='top', bbox=dict(facecolor='white', alpha=0.5))

    def export_plots(self):
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading('G-Levels and PSD Plots', 0)

        for fig in self.glevel_figs:
            img_path = '/tmp/glevel_plot.png'
            fig.savefig(img_path)
            doc.add_paragraph('G-Level Plot')
            doc.add_picture(img_path, width=Inches(6))

        for fig in self.psd_figs:
            img_path = '/tmp/psd_plot.png'
            fig.savefig(img_path)
            doc.add_paragraph('PSD Plot')
            doc.add_picture(img_path, width=Inches(6))

        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx")])
        if save_path:
            doc.save(save_path)
            messagebox.showinfo("Export Successful", "Plots exported successfully.")

if __name__ == "__main__":
    app = GLevelPSDApp()
    app.mainloop()
